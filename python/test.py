import win32com.client
from docx import Document
import os
import openpyxl
import re
from datetime import datetime


def convert_doc_to_docx(doc_path):
    """
    将 .doc 文件转换为 .docx
    """
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False  # 设为 True 可查看 Word 运行状态

    try:
        doc = word.Documents.Open(doc_path)
        docx_path = doc_path + "x"  # 直接在原路径后加 "x"
        doc.SaveAs(docx_path, FileFormat=16)  # 16 表示 docx 格式
        doc.Close()
        return docx_path
    except Exception as e:
        print("转换失败:", e)
        return None
    finally:
        word.Quit()

from docx import Document
import re
from datetime import datetime

def extract_meeting_datetime(docx_path):
    """从表格中提取会议时间（格式：yyyyMMdd）"""
    doc = Document(docx_path)
    
    # 正则表达式，匹配形如 "2025年3月7日上午9:00" 或 "2025年3月7日9:00" 的日期时间
    meeting_time_pattern = r'(\d{4})年(\d{1,2})月(\d{1,2})日'
    
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if cell.text.strip() == '会议时间':
                    # 遍历后续单元格，寻找会议时间
                    for j in range(i + 1, len(row.cells)):
                        meeting_time = row.cells[j].text.strip()
                        
                        # 使用正则表达式查找匹配的时间
                        match = re.search(meeting_time_pattern, meeting_time)
                        if match:
                            year, month, day = match.groups()
                            # 拼接日期并解析
                            meeting_date_str = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                            try:
                                # 解析并格式化日期为 yyyyMMdd
                                meeting_date = datetime.strptime(meeting_date_str, '%Y-%m-%d')
                                return meeting_date.strftime('%Y%m%d')
                            except ValueError:
                                pass  # 若解析失败则继续搜索下一个单元格
    return None  # 若未找到有效的会议时间

def extract_meeting_time(docx_path):
    """
    从 Word 文档中提取会议时间（格式：第X次党委会）
    """
    doc = Document(docx_path)
    pattern = r'(第\d+次党委会)'

    # 检查所有段落
    for para in doc.paragraphs:
        text = para.text.strip()
        match = re.search(pattern, text)
        if match:
            return match.group(1)

    # 检查表格中的每个单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                match = re.search(pattern, text)
                if match:
                    return match.group(1)

    return None

def extract_meeting_agenda_from_table(docx_path):
    """
    从 Word 表格中提取会议议程并去重（保留顺序）
    """
    doc = Document(docx_path)
    agenda_content = []

    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            if "会议议程" in row_texts[0]:
                agenda_content.append(row_texts[1])

    return agenda_content if agenda_content else ["未找到会议议程"]

def convert_agenda_content(content):
    result = []
    count = 1  # 序号计数器

    # 中文大写数字
    chinese_numbers = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
                       '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']

    def to_chinese_number(n):
        """将数字转换为中文大写数字序号"""
        if n <= 20:
            return chinese_numbers[n - 1]
        else:
            # 对大于20的数字进行处理，按十进制分割
            tens = n // 10
            ones = n % 10
            tens_part = chinese_numbers[tens - 1] + '十' if tens > 1 else '十'
            ones_part = chinese_numbers[ones - 1] if ones > 0 else ''
            return tens_part + ones_part

    for line in content.split('\n'):
        line = line.strip()  # 去除首尾空格
        if '（汇报人：' in line or '（传达人：' in line:  # 只处理包含汇报人的行
            # 去除行首的数字编号（如1.、2.）及可能存在的空格
            processed_line = re.sub(r'^\d+\.\s*', '', line)
            # 添加大写序号
            number = to_chinese_number(count)
            processed_line = f"{number}、{processed_line}"
            result.append(processed_line)
            count += 1

    return result

def write_agenda_to_excel(excel_path, sheet_name, agenda_list, meeting_time):
    """
    将会议时间和议程追加写入 Excel 的空白行（时间在 B 列，议程在 C 列）
    """
    wb = openpyxl.load_workbook(excel_path)
    if sheet_name not in wb.sheetnames:
        print(f"错误: Excel 文件中未找到工作表 '{sheet_name}'")
        return False

    ws = wb[sheet_name]

    # 查找C列最后一个非空行（从下往上找）
    max_row = ws.max_row  # 获取当前工作表最大行号
    start_row = 2  # 默认起始行

    # 如果发现C列已经有数据，找到最后一行数据的下一行
    for row_num in range(max_row, 1, -1):
        if ws.cell(row=row_num, column=3).value is not None:  # column=3 对应C列
            start_row = row_num + 1
            break

    # 写入数据（保持时间和议程的对应关系）
    for i, item in enumerate(agenda_list):
        current_row = start_row + i
        ws[f"B{current_row}"] = meeting_time  # B列写入时间
        ws[f"C{current_row}"] = item          # C列写入议程内容

    wb.save(excel_path)
    wb.close()
    return True


# 文件路径（请修改为你的实际路径）
doc_path = r"C:/Users/ashin/Nextcloud/同步盘/ashin/Desktop/0.赣州职业技术学院2025年第2次党委会会议方案(2).doc"
excel_path = r"C:/Users/ashin/Nextcloud/同步盘/ashin/Desktop/议题记录.xlsx"
sheet_name = "党委会-2025"

# 1. 判断文件类型
file_extension = os.path.splitext(doc_path)[1].lower()

if file_extension == '.doc':
    # 转换 DOC 到 DOCX
    docx_path = convert_doc_to_docx(doc_path)

    if docx_path:
        # 2. 提取会议时间（优先从内容提取，失败则从文件名提取）
        meeting_time = extract_meeting_datetime(docx_path) + " " + extract_meeting_time(docx_path)
        if not meeting_time:
            meeting_time = "未知时间"
            print("警告：未检测到会议时间")

        # 3. 提取会议议程
        agenda = extract_meeting_agenda_from_table(docx_path)
        if agenda[0] == "未找到会议议程":
            print(agenda[0])
        else:
            result = convert_agenda_content(agenda[0])

            # 4. 写入 Excel
            if write_agenda_to_excel(excel_path, sheet_name, result, meeting_time):
                print("会议时间和议程已成功写入 Excel 文件！")

        # 删除转换后的 docx 文件（可选）
        os.remove(docx_path)

elif file_extension == '.docx':
    # 如果是 docx 文件，直接提取议程并写入 Excel
    docx_path = doc_path

    # 提取会议时间（优先从内容提取，失败则从文件名提取）
    meeting_time = extract_meeting_datetime(docx_path) + " " + extract_meeting_time(docx_path)
    if not meeting_time:
        meeting_time = "未知时间"
        print("警告：未检测到会议时间")

    # 提取会议议程
    agenda = extract_meeting_agenda_from_table(docx_path)
    if agenda[0] == "未找到会议议程":
        print(agenda[0])
    else:
        result = convert_agenda_content(agenda[0])

        # 写入 Excel
        if write_agenda_to_excel(excel_path, sheet_name, result, meeting_time):
            print("会议时间和议程已成功写入 Excel 文件！")

else:
    print("文件格式不支持。")
