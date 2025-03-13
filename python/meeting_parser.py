import win32com.client
from docx import Document
import os
import openpyxl
import re
from datetime import datetime
import logging

# 配置详细日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("meeting_parser.log", encoding='utf-8'),
        logging.StreamHandler()
    ]
)

def num_to_chinese(n):
    """将数字转换为中文表示（支持1-99）"""
    logging.debug(f"开始数字转换: {n}")
    chinese_digits = ['', '一', '二', '三', '四', '五', '六', '七', '八', '九']
    try:
        if n < 1:
            return ''
        if n < 10:
            result = chinese_digits[n]
        elif n < 20:
            result = '十' + chinese_digits[n % 10] if n % 10 != 0 else '十'
        elif n < 100:
            tens = chinese_digits[n // 10]
            remainder = n % 10
            result = f"{tens}十{chinese_digits[remainder]}" if remainder !=0 else f"{tens}十"
        else:
            result = str(n)
        logging.debug(f"数字转换结果: {n} -> {result}")
        return result
    except Exception as e:
        logging.error(f"数字转换异常: {str(e)}")
        return str(n)

def convert_doc_to_docx(doc_path):
    """将 .doc 文件转换为 .docx"""
    logging.info(f"开始文件转换: {doc_path}")
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc = word.Documents.Open(doc_path)
        docx_path = os.path.splitext(doc_path)[0] + ".docx"
        logging.debug(f"打开文档成功，开始保存为: {docx_path}")
        
        doc.SaveAs(docx_path, FileFormat=16)
        doc.Close()
        word.Quit()
        logging.info(f"文件转换完成，新路径: {docx_path}")
        return docx_path
    except Exception as e:
        logging.error(f"文件转换失败 | 路径: {doc_path} | 错误: {str(e)}", exc_info=True)
        return None
    finally:
        if 'word' in locals():
            logging.debug("正在释放Word进程资源")
            word.Quit()

def extract_meeting_info(docx_path):
    """从文档中提取结构化会议信息"""
    logging.info(f"开始提取会议信息: {docx_path}")
    doc = Document(docx_path)
    meeting_info = {'date': None, 'meeting_num': None}

    # 表格搜索部分
    logging.debug("开始表格搜索")
    for table_idx, table in enumerate(doc.tables):
        logging.debug(f"正在处理第{table_idx+1}个表格")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                logging.debug(f"表格{table_idx+1}-行{row_idx+1}-格{cell_idx+1}: {text}")

                # 日期提取
                if not meeting_info['date'] and "会议时间" in text:
                    logging.debug("发现日期关键词，检查后续单元格")
                    for next_cell in row.cells[cell_idx+1:]:
                        match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', next_cell.text)
                        if match:
                            year, month, day = match.groups()
                            meeting_info['date'] = f"{year}{month.zfill(2)}{day.zfill(2)}"
                            logging.info(f"从表格中提取到日期: {meeting_info['date']}")
                            break

                # 会议次数提取
                if not meeting_info['meeting_num']:
                    match = re.search(r'(第\d+次党委会)', text)
                    if match:
                        meeting_info['meeting_num'] = match.group(1)
                        logging.info(f"从表格中提取到会议次数: {meeting_info['meeting_num']}")

    # 段落补充搜索
    logging.debug("开始段落搜索")
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            logging.debug(f"段落{para_idx+1}: {text}")

        # 补充日期提取
        if not meeting_info['date']:
            date_match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', text)
            if date_match:
                year, month, day = date_match.groups()
                meeting_info['date'] = f"{year}{month.zfill(2)}{day.zfill(2)}"
                logging.info(f"从段落中补充提取到日期: {meeting_info['date']}")

        # 补充会议次数提取
        if not meeting_info['meeting_num']:
            num_match = re.search(r'(第\d+次党委会)', text)
            if num_match:
                meeting_info['meeting_num'] = num_match.group(1)
                logging.info(f"从段落中补充提取到会议次数: {meeting_info['meeting_num']}")

    logging.info(f"最终会议信息: {meeting_info}")
    return meeting_info

def extract_info_from_filename(filename):
    """从文件名提取会议次数"""
    logging.debug(f"开始文件名解析: {filename}")
    match = re.search(r'第(\d+)次党委会', filename)
    if match:
        result = f"第{match.group(1)}次党委会"
        logging.info(f"从文件名提取到会议次数: {result}")
        return result
    logging.warning("文件名中未找到有效会议次数")
    return None

def parse_agenda_content(raw_text):
    """处理原始议程文本"""
    logging.info("开始解析议程内容")
    agenda_items = []
    current_num = 1
    total_lines = 0
    valid_items = 0

    for line in raw_text.split('\n'):
        total_lines += 1
        line = line.strip()
        if not line:
            logging.debug("跳过空行")
            continue

        clean_line = re.sub(r'^[\d、.]+', '', line)
        reporter_match = re.search(r'（(汇报人|传达人)：(.+?)）', clean_line)
        
        if not reporter_match:
            logging.debug(f"跳过无汇报人条目: {line}")
            continue

        valid_items += 1
        cn_num = num_to_chinese(current_num)
        agenda_item = {
            'index': current_num,
            'cn_index': f"{cn_num}、",
            'content': clean_line
        }
        logging.debug(f"解析到有效议程项: {agenda_item}")
        agenda_items.append(agenda_item)
        current_num += 1

    logging.info(f"议程解析完成 | 总行数: {total_lines} | 有效条目: {valid_items}")
    return agenda_items

def get_agenda_text(docx_path):
    """从文档表格中提取议程原始文本"""
    logging.info(f"开始提取议程原始文本: {docx_path}")
    doc = Document(docx_path)
    
    for table_idx, table in enumerate(doc.tables):
        logging.debug(f"正在扫描第{table_idx+1}个表格")
        for row_idx, row in enumerate(table.rows):
            if row.cells[0].text.strip() == "会议议程":
                agenda_cell = row.cells[1].text.strip()
                logging.info(f"找到议程单元格: {agenda_cell[:50]}...")  # 防止日志过长
                return agenda_cell
    logging.warning("未找到会议议程表格")
    return None

def write_to_excel(excel_path, sheet_name, data):
    """将数据写入Excel"""
    logging.info(f"开始写入Excel | 文件: {excel_path} | 工作表: {sheet_name}")
    try:
        wb = openpyxl.load_workbook(excel_path)
        logging.debug("Excel文件加载成功")
        
        if sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            logging.info(f"使用现有工作表: {sheet_name}")
        else:
            ws = wb.create_sheet(sheet_name)
            logging.info(f"创建新工作表: {sheet_name}")

        # 查找最后一行
        last_row = 1
        while ws.cell(row=last_row+1, column=3).value is not None:
            last_row += 1
        logging.debug(f"检测到最后有效行: {last_row}")

        # 写入数据
        for idx, item in enumerate(data['agenda'], start=1):
            row_num = last_row + idx
            ws[f'B{row_num}'] = data['date']
            ws[f'C{row_num}'] = f"{item['cn_index']}{item['content']}"
            logging.debug(f"写入行{row_num}: {item['content'][:30]}...")

        wb.save(excel_path)
        logging.info(f"成功写入{len(data['agenda'])}条数据 | 起始行: {last_row+1}")
        return True
    except Exception as e:
        logging.error(f"Excel写入失败 | 路径: {excel_path} | 错误: {str(e)}", exc_info=True)
        return False

def main(input_path, excel_path, sheet_name):
    """主处理流程"""
    logging.info("="*50)
    logging.info(f"开始处理文件: {input_path}")
    
    try:
        # 文件类型处理
        base_path, ext = os.path.splitext(input_path)
        if ext == '.doc':
            logging.info("检测到.doc文件，启动格式转换")
            docx_path = convert_doc_to_docx(input_path)
            if not docx_path:
                logging.error("文件转换失败，终止流程")
                return
        else:
            docx_path = input_path
            logging.info("已是.docx格式，跳过转换")

        # 信息提取
        logging.info("开始信息提取流程")
        meeting_info = extract_meeting_info(docx_path)
        filename_info = extract_info_from_filename(os.path.basename(input_path))
        
        # 信息合并
        final_date = meeting_info['date'] or datetime.now().strftime("%Y%m%d")
        final_num = meeting_info['meeting_num'] or filename_info or "第X次党委会"
        logging.info(f"最终会议信息 | 日期: {final_date} | 次数: {final_num}")

        # 获取议程
        logging.info("开始提取议程内容")
        raw_agenda = get_agenda_text(docx_path)
        if not raw_agenda:
            logging.error("未找到会议议程内容，终止流程")
            return

        # 处理数据
        logging.info("开始解析议程内容")
        processed_data = {
            'date': f"{final_date} {final_num}",
            'agenda': parse_agenda_content(raw_agenda)
        }
        logging.debug(f"处理后的数据结构: {processed_data.keys()}")

        # 写入Excel
        logging.info("开始写入Excel")
        if write_to_excel(excel_path, sheet_name, processed_data):
            logging.info("处理流程完成")
        else:
            logging.error("写入Excel失败")

    except Exception as e:
        logging.error(f"主流程异常: {str(e)}", exc_info=True)
    finally:
        if ext == '.doc' and os.path.exists(docx_path):
            logging.info(f"开始清理临时文件: {docx_path}")
            try:
                os.remove(docx_path)
                logging.info("临时文件清理成功")
            except Exception as e:
                logging.error(f"文件清理失败: {str(e)}")
        logging.info("处理流程结束\n")

if __name__ == "__main__":
    # 配置参数
    config = {
        'input_path': r"C:/Users/ashin/Nextcloud/同步盘/ashin/Desktop/0.赣州职业技术学院2025年第2次党委会会议方案(2).doc",
        'excel_path': r"C:/Users/ashin/Nextcloud/同步盘/ashin/Desktop/议题记录.xlsx",
        'sheet_name': "党委会-2025"
    }   
    
    main(**config)
